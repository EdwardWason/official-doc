"""
Official Doc Formatter - 公文格式转换核心模块
基于 GB/T 9704-2012 标准，将 Markdown 转换为党政机关公文格式
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

__version__ = "1.0.0"


def set_line_spacing(paragraph, spacing_pt=26):
    """设置固定行间距 26 磅"""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:line'), str(int(spacing_pt * 20)))
    spacing.set(qn('w:lineRule'), 'exact')
    pPr.append(spacing)


def set_first_line_indent(paragraph, chars=3):
    """设置首行缩进 3 字符"""
    pPr = paragraph._p.get_or_add_pPr()
    indent = OxmlElement('w:ind')
    indent.set(qn('w:firstLine'), str(int(chars * 20 * 12)))
    pPr.append(indent)


def set_font(run, font_name, font_size, bold=False):
    """设置字体样式"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_page_number(doc):
    """添加页码 - 居中显示，样式：— 1 —"""
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.text = '— '
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar)
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        run._element.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar2)
        paragraph.add_run(' —')
        for run in paragraph.runs:
            run.font.name = '宋体'
            run.font.size = Pt(14)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def remove_all_spaces(text):
    """删除所有空格"""
    text = text.replace(' ', '')
    return text


def convert_quotes(text):
    """转换英文引号为中文引号"""
    text = re.sub(r'"([^"]+)"', r'"\1"', text)
    text = re.sub(r"''([^']+)'", r"'\1'", text)
    text = re.sub(r"'([^']+)'", r"'\1'", text)
    text = re.sub(r'""', r'""', text)
    text = re.sub(r"''", r"''", text)
    return text


def clean_text(text):
    """清理文本：删除空格 + 转换引号"""
    text = remove_all_spaces(text)
    text = convert_quotes(text)
    return text


def parse_bold_runs(text, paragraph, font_name, font_size):
    """解析加粗文本，支持 **text** 语法"""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            set_font(run, font_name, font_size, bold=True)
        elif part:
            run = paragraph.add_run(part)
            set_font(run, font_name, font_size, bold=False)


def md_to_docx(md_content, output_path):
    """
    将 Markdown 内容转换为符合 GB/T 9704-2012 标准的 Word 文档

    参数:
        md_content (str): Markdown 格式的文本内容
        output_path (str): 输出 Word 文件路径

    返回:
        bool: 转换是否成功

    示例:
        >>> md_content = "# 标题\\n\\n这是正文内容"
        >>> md_to_docx(md_content, "output.docx")
        True
    """
    try:
        doc = Document()

        section = doc.sections[0]
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)

        lines = md_content.split('\n')

        level1_counter = 0
        level2_counter = 0
        level3_counter = 0
        level4_counter = 0

        chinese_nums = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
                        '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十']

        for line in lines:
            if not line.strip():
                continue

            if line.startswith('# '):
                level1_counter += 1
                level2_counter = 0
                level3_counter = 0
                level4_counter = 0

                title_text = line[2:].strip()
                title_text = clean_text(title_text)
                if level1_counter <= len(chinese_nums):
                    title_text = f"{chinese_nums[level1_counter-1]}、{title_text}"
                else:
                    title_text = f"{level1_counter}、{title_text}"

                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_line_spacing(paragraph, 26)
                run = paragraph.add_run(title_text)
                set_font(run, '方正小标宋简体', 22, bold=False)

            elif line.startswith('## '):
                level2_counter += 1
                level3_counter = 0
                level4_counter = 0

                heading_text = line[3:].strip()
                heading_text = clean_text(heading_text)
                if level2_counter <= len(chinese_nums):
                    heading_text = f"{chinese_nums[level2_counter-1]}、{heading_text}"
                else:
                    heading_text = f"{level2_counter}、{heading_text}"

                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_line_spacing(paragraph, 26)
                set_first_line_indent(paragraph, 3)
                run = paragraph.add_run(heading_text)
                set_font(run, '黑体', 16, bold=False)

            elif line.startswith('### '):
                level3_counter += 1
                level4_counter = 0

                heading_text = line[4:].strip()
                heading_text = clean_text(heading_text)
                heading_text = f"（{chinese_nums[level3_counter-1]}）{heading_text}"

                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_line_spacing(paragraph, 26)
                set_first_line_indent(paragraph, 3)
                run = paragraph.add_run(heading_text)
                set_font(run, '楷体_GB2312', 16, bold=True)

            elif line.startswith('#### '):
                level4_counter += 1

                heading_text = line[5:].strip()
                heading_text = clean_text(heading_text)
                heading_text = f"{level4_counter}.{heading_text}"

                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_line_spacing(paragraph, 26)
                set_first_line_indent(paragraph, 3)
                run = paragraph.add_run(heading_text)
                set_font(run, '仿宋_GB2312', 16, bold=True)

            elif line.startswith('> '):
                continue

            elif line.startswith('---'):
                continue

            elif line.startswith('![]('):
                continue

            elif line.startswith('- ') or line.startswith('* '):
                list_item = line[2:].strip()
                list_item = clean_text(list_item)
                paragraph = doc.add_paragraph()
                set_line_spacing(paragraph, 26)
                set_first_line_indent(paragraph, 3)
                parse_bold_runs(list_item, paragraph, '仿宋_GB2312', 16)

            elif re.match(r'^\d+\.', line):
                list_item = re.sub(r'^\d+\.', '', line).strip()
                list_item = clean_text(list_item)
                paragraph = doc.add_paragraph()
                set_line_spacing(paragraph, 26)
                set_first_line_indent(paragraph, 3)
                parse_bold_runs(list_item, paragraph, '仿宋_GB2312', 16)

            else:
                clean_line = line.strip()
                clean_line = clean_text(clean_line)
                if clean_line:
                    paragraph = doc.add_paragraph()
                    set_line_spacing(paragraph, 26)
                    set_first_line_indent(paragraph, 3)
                    parse_bold_runs(clean_line, paragraph, '仿宋_GB2312', 16)

        add_page_number(doc)
        doc.save(output_path)
        return True

    except Exception as e:
        print(f"转换失败: {str(e)}")
        return False